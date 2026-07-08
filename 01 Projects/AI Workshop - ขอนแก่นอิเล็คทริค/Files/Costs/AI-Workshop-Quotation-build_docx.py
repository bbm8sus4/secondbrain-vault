#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE   = RGBColor(0x12,0x46,0x7e)
GRAY   = RGBColor(0x66,0x66,0x66)
RED    = RGBColor(0xc0,0x39,0x2b)
GREEN  = RGBColor(0x2e,0x7d,0x32)
BLACK  = RGBColor(0x1a,0x1a,0x1a)
HDR_BG   = "E9EDF2"
SEC_BG   = "EAF1F9"
TOT_BG   = "F6F8FA"
FONT = "Sarabun"

doc = Document()
# page setup
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21)
sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
USABLE = 21 - 2.0 - 2.0  # 17.0 cm

# base font
st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:cs'), FONT)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def shade(cell, hexclr):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexclr)
    tcPr.append(sh)

def set_cell_borders(cell, color="C9D0D9", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz); e.set(qn('w:space'),'0'); e.set(qn('w:color'),color)
        borders.append(e)
    tcPr.append(borders)

def fixed_layout(table):
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed')
    tblPr.append(layout)

def run(p, text, size=10.5, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r._element.rPr.rFonts.set(qn('w:cs'), FONT)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if color is not None: r.font.color.rgb = color
    # complex-script bold/size
    rPr = r._element.rPr
    if bold:
        b = OxmlElement('w:bCs'); rPr.append(b)
    szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), str(int(size*2))); rPr.append(szcs)
    return r

def para(container, align=None, space_after=2, space_before=0):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before); pf.line_spacing = 1.05
    if align is not None: p.alignment = align
    return p

def set_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

# ---------- Header ----------
p = para(doc, WD_ALIGN_PARAGRAPH.RIGHT, space_after=0); run(p,"THUNDER",15,True,BLUE)
p = para(doc, WD_ALIGN_PARAGRAPH.RIGHT, space_after=6); run(p,"SOLUTION · AUTOMATION · AI",7.5,False,GRAY)

p = para(doc, space_after=8); run(p,"Date: 8 กรกฎาคม 2026   ·   เลขที่ใบเสนอราคา: QT-KKAI-20260708-02",10.5,True)

# ---------- Parties (2-col borderless) ----------
pt = doc.add_table(rows=1, cols=2); fixed_layout(pt); set_widths(pt,[8.5,8.5])
def party(cell, role, name, lines, tail):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); run(p, role, 10, True, BLUE)
    p = para(cell, space_after=1); run(p, name, 10.5, True)
    for ln in lines:
        p = para(cell, space_after=0)
        if "[" in ln:
            pre, ph = ln.split("[",1); run(p,pre,9.5,False,GRAY); run(p,"["+ph,9.5,True,RED)
        else:
            run(p, ln, 9.5, False, GRAY)
    p = para(cell, space_after=1); run(p, tail, 9.5)
party(pt.rows[0].cells[0], "BETWEEN (ผู้ว่าจ้าง / CLIENT):",
      "บริษัท ขอนแก่น อีเลคทริค เทคโนโลยี เอนจิเนียริ่ง จำกัด",
      ["KHONKAEN ELECTRIC TECHNOLOGY ENGINEERING CO., LTD.",
       "888/216 หมู่บ้าน ศุภาลัย การ์เด้นวิลล์ หมู่ที่ 14",
       "ตำบลบ้านเป็ด อำเภอเมืองขอนแก่น จังหวัดขอนแก่น 40000",
       "ผู้ติดต่อ: คุณป๊อบ",
       "Tax ID: [เลขผู้เสียภาษีลูกค้า]"],
      "(เรียกว่า “ผู้ว่าจ้าง”)")
party(pt.rows[0].cells[1], "AND (ผู้ให้บริการ / PROVIDER):",
      "บริษัท ธันเดอร์ โซลูชั่น จำกัด",
      ["THUNDER SOLUTION CO., LTD.",
       "629 หมู่ที่ 6 ตำบลบ้านเป็ด",
       "อำเภอเมืองขอนแก่น จังหวัดขอนแก่น 40000",
       "Tel: +66 (0) 9 372 8583 (K.Ghee)",
       "Email: adisorn.n@thunder.in.th",
       "Tax ID: 0465566000017"],
      "(เรียกว่า “Thunder”)")

# ---------- H1 ----------
def h1(text):
    p = para(doc, space_before=10, space_after=4); run(p, text, 16, True, BLUE)
    # bottom border
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'2'); bottom.set(qn('w:color'),'B9C4D2')
    pbdr.append(bottom); pPr.append(pbdr)
def h2(text):
    p = para(doc, space_before=10, space_after=3); run(p, text, 12.5, True, BLUE)

h1("QUOTATION OF SERVICE")

# ---------- Tasks table ----------
def hdr_cell(cell, text, align=None, bg=HDR_BG):
    shade(cell, bg); set_cell_borders(cell)
    cell.paragraphs[0].text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(1)
    if align is not None: p.alignment=align
    run(p, text, 10.5, True)

tt = doc.add_table(rows=2, cols=2); fixed_layout(tt); set_widths(tt,[8.5,8.5])
hdr_cell(tt.rows[0].cells[0],"Tasks (ขอบเขตงาน)")
hdr_cell(tt.rows[0].cells[1],"Estimated Duration (ระยะเวลา)")
for c,txt in zip(tt.rows[1].cells,
   ["อบรมเชิงปฏิบัติการ “AI Workshop” การใช้ Claude AI กับงานจริง สำหรับพนักงาน 15 คน (คละแผนก)",
    "อบรม 1 วัน (09:00–16:30 น.) · รวมงานเตรียม Needs Analysis + Curriculum ก่อนอบรม"]):
    set_cell_borders(c); c.paragraphs[0].text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); run(p,txt,10)

para(doc, space_after=2)

# ---------- Services table ----------
SW=[9.5,2.3,1.5,3.2]
sv = doc.add_table(rows=2, cols=4); fixed_layout(sv); set_widths(sv,SW)
hdr_cell(sv.rows[0].cells[0],"Services (รายละเอียดบริการ)")
hdr_cell(sv.rows[0].cells[1],"Cost/Unit",WD_ALIGN_PARAGRAPH.CENTER)
hdr_cell(sv.rows[0].cells[2],"Unit",WD_ALIGN_PARAGRAPH.CENTER)
hdr_cell(sv.rows[0].cells[3],"Cost (THB)",WD_ALIGN_PARAGRAPH.CENTER)
# section row (merge)
secrow = sv.rows[1]
a = secrow.cells[0].merge(secrow.cells[1]).merge(secrow.cells[2]).merge(secrow.cells[3])
shade(a, SEC_BG); set_cell_borders(a)
a.paragraphs[0].text=""; p=a.paragraphs[0]; p.paragraph_format.space_after=Pt(1)
run(p,"Thunder | AI Workshop Service | อบรม 1 วัน (15 ท่าน)",10.5,True,BLUE)
# body row
body = sv.add_row(); set_widths(sv,SW)
bc = body.cells[0]; set_cell_borders(bc); bc.paragraphs[0].text=""
def group(cell, title, items):
    p = para(cell, space_before=3, space_after=1); run(p, "● "+title, 10.5, True)
    for it in items:
        p = para(cell, space_after=0); p.paragraph_format.left_indent=Cm(0.5)
        run(p, "•  ", 10)
        # bold lead before ':' if present with (Track)
        if it.startswith("**"):
            _, rest = it.split("**",1); lead, rest2 = rest.split("**",1)
            run(p, lead, 10, True); run(p, rest2, 10)
        else:
            run(p, it, 10)
group(bc,"SETUP & PREPARATION (One Time)",
     ["Needs Analysis: วิเคราะห์แบบสำรวจผู้เรียน 11 คน + จัดกลุ่ม tiered track ตามความพร้อม",
      "Curriculum Design: ออกแบบ agenda นาทีต่อนาที (Core / Advanced / Demo)",
      "เตรียมบัญชี/สภาพแวดล้อม demo + Fallback pack (กันเน็ตล่ม)"])
group(bc,"TRAINING DELIVERY (On-site 1 Day · 15 pax)",
     ["**Core Track**  (ทุกคน · browser-only): Claude พื้นฐาน, Prompt Foundation, งานคอนเทนต์/สรุป/แปล/วิเคราะห์, ทำ Landing Page จาก template",
      "**Advanced Track**  (กลุ่มพร้อม): Cowork Basic, Claude CLI Basic — hands-on แบบ buddy",
      "**Demo Track**  (ดูบนจอ): Claude Code / Warp overview, Deploy / GitHub / Database",
      "Capstone + Showcase ชิ้นงานจริงของผู้เรียน"])
group(bc,"MATERIALS & DELIVERABLES",
     ["สไลด์ประกอบการอบรม (soft copy)",
      "Lab Exercises + Prompt Template Pack",
      "Learner Dashboard 2 แท็บ (mobile) + Cheat Sheet",
      "Offline / Fallback Pack"])
group(bc,"SUPPORT & REPORTING",
     ["แบบประเมินหลังอบรม + สรุปผลประเมิน",
      "Runbook หน้างาน + ระบบ Buddy"])
for idx,val in [(1,"2,499.00"),(2,"15"),(3,"37,485.00")]:
    c = body.cells[idx]; set_cell_borders(c); c.paragraphs[0].text=""
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx!=2 else WD_ALIGN_PARAGRAPH.CENTER
    run(p, val, 10.5, True)

para(doc, space_after=2)

# ---------- Totals (right-aligned) ----------
tot = doc.add_table(rows=3, cols=2); fixed_layout(tot); tot.alignment = WD_TABLE_ALIGNMENT.RIGHT
set_widths(tot,[4.0,2.8])
rows = [("Total (ก่อนภาษี)","37,485.00",TOT_BG,False),
        ("VAT 7%","2,623.95",TOT_BG,False),
        ("Grand Total","40,108.95",SEC_BG,True)]
for i,(lbl,val,bg,grand) in enumerate(rows):
    lc,vc = tot.rows[i].cells
    for c in (lc,vc): shade(c,bg); set_cell_borders(c)
    lc.paragraphs[0].text=""; p=lc.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(1)
    run(p,lbl,12 if grand else 10.5,True, BLUE if grand else None)
    vc.paragraphs[0].text=""; p=vc.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(1)
    run(p,val,12 if grand else 10.5,True, BLUE if grand else None)

# ---------- KPI ----------
h2("Estimated Outcomes (KPI)")
KW=[1.8,2.0,7.5,5.0]
kpi = doc.add_table(rows=1, cols=4); fixed_layout(kpi); set_widths(kpi,KW)
for c,t in zip(kpi.rows[0].cells,["Track","ผู้เรียน","ผลลัพธ์ที่คาดหวัง","วิธีวัดผล"]):
    hdr_cell(c,t)
kdata=[("Core","ทุกคน (15)","ใช้ Claude ทำงานจริงได้ — เขียนคอนเทนต์ / สรุป / แปล / วิเคราะห์เบื้องต้น + ทำ Landing Page จาก template","แบบประเมินหลังอบรม + ชิ้นงาน Capstone"),
       ("Advanced","กลุ่มพร้อม","ใช้ Cowork + Claude CLI Basic ได้ด้วยตนเอง (มี buddy)","Hands-on ผ่าน + buddy sign-off"),
       ("Demo","ทุกคน (ดู)","เข้าใจภาพรวม Deploy / GitHub / Database","Q&A + ความมั่นใจเพิ่มขึ้น (survey ก่อน/หลัง)")]
for tr,pu,res,how in kdata:
    r = kpi.add_row(); set_widths(kpi,KW)
    for c,txt,bold in zip(r.cells,[tr,pu,res,how],[True,False,False,False]):
        set_cell_borders(c); c.paragraphs[0].text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); run(p,txt,10,bold)

# ---------- Conditions ----------
h2("Conditions & Remarks (เงื่อนไข)")
conds=[("ESTIMATE"," ตัวเลขและผลลัพธ์เป็นการประมาณการตามข้อมูลผู้เรียนที่สำรวจได้"),
       (None,"กรุณายืนยันใบเสนอราคาภายใน 15 วัน นับจากวันที่ในเอกสาร มิฉะนั้นถือว่าสิ้นผล"),
       (None,"ราคานี้รวมค่าวิทยากร (facilitation) + วัสดุการอบรมแบบ soft copy สำหรับผู้เรียน 15 ท่าน"),
       ("ไม่รวม:"," ค่า subscription ของ Claude หรือเครื่องมือภายนอกที่ผู้เรียนใช้เอง · ค่า cloud/deploy/database หากใช้งานจริงนอกเหนือจาก demo · ค่าเดินทาง/ที่พักกรณีจัดนอกจังหวัดขอนแก่น · ค่าพิมพ์เอกสารเป็นกระดาษ"),
       (None,"สถานที่ อาหาร อินเทอร์เน็ต และอุปกรณ์ (โน้ตบุ๊กผู้เรียน) ผู้ว่าจ้างเป็นผู้จัดเตรียม"),
       (None,"ผู้เรียนควรมีโน้ตบุ๊ก 1 เครื่อง/คน และบัญชี Claude (Core ใช้งานผ่าน browser)"),
       ("ห้าม","นำข้อมูลลับ/ข้อมูลส่วนบุคคล/รหัสผ่านจริงเข้าสู่ AI ระหว่างการอบรม เพื่อความปลอดภัยของข้อมูล"),
       (None,"ค่าบริการชำระเต็มจำนวนก่อนเริ่มงาน และเป็นจำนวนที่ไม่สามารถขอคืนได้ (non-refundable) เมื่อ Thunder เริ่มงานเตรียมแล้ว")]
for lead,txt in conds:
    p = para(doc, space_after=1); p.paragraph_format.left_indent=Cm(0.5)
    run(p,"•  ",10)
    if lead=="ESTIMATE": run(p," ESTIMATE ",9,True,GREEN)
    elif lead: run(p,lead,10,True, RED if lead=="ห้าม" else None)
    run(p,txt,10)

# ---------- Payment ----------
h2("Payment (การชำระเงิน)")
p = para(doc, space_after=1); run(p,"ชำระเต็มจำนวนครั้งเดียว: ",10.5); run(p,"40,108.95 THB",11,True,RED); run(p,"  (37,485.00 + VAT 7% = 2,623.95)",10.5)
p = para(doc, space_after=2); run(p,"ออกใบวางบิลเมื่อยืนยันงาน · ครบกำหนดชำระก่อนวันอบรม (อย่างน้อย 7 วันก่อนวันอบรม)",10)

# ---------- Transfer ----------
h2("Transfer Details (รายละเอียดการโอนเงิน)")
trf=[("ชื่อบัญชี:","บริษัท ธันเดอร์ โซลูชั่น จำกัด ","[ยืนยันชื่อบัญชีจริง]"),
     ("เลขที่บัญชี:","","[กรอกเลขบัญชี Thunder]"),
     ("ธนาคาร / สาขา:","","[ชื่อธนาคาร / สาขา]  (ออมทรัพย์)")]
for lbl,val,ph in trf:
    p = para(doc, space_after=1); run(p,lbl+"  ",10,True); run(p,val,10)
    if ph: run(p,ph,10,True,RED)
p = para(doc, space_after=2); run(p,"ส่งหลักฐานการโอนมาที่ Email: adisorn.n@thunder.in.th (โทร. 09 372 8583 · K.Ghee)",9.5,False,GRAY)

# ---------- Signatures ----------
h2("Signatures (ลงนาม)")
sg = doc.add_table(rows=1, cols=2); fixed_layout(sg); set_widths(sg,[8.5,8.5])
def sig(cell, co, lines):
    cell.paragraphs[0].text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); run(p,co,10.5,True)
    para(cell, space_after=10)
    for ln in lines:
        p=para(cell, space_after=2); run(p,ln,10)
sig(sg.rows[0].cells[0],"บริษัท ขอนแก่น อีเลคทริค เทคโนโลยี เอนจิเนียริ่ง จำกัด (ผู้ว่าจ้าง)",
    ["ลายเซ็น: ____________________________","วันที่: ____________________________","ชื่อ: คุณป๊อบ ______________________","ตำแหน่ง: ____________________________"])
sig(sg.rows[0].cells[1],"บริษัท ธันเดอร์ โซลูชั่น จำกัด (ผู้ให้บริการ)",
    ["ลายเซ็น: ____________________________","วันที่: 8 กรกฎาคม 2026","ชื่อ: คุณอดิศร (K.Ghee) __________","ตำแหน่ง: ผู้มีอำนาจลงนาม"])

import sys
out = sys.argv[1] if len(sys.argv)>1 else "out.docx"
doc.save(out)
print("saved", out)
